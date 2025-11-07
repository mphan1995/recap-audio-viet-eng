import time , torch , warnings , json ,os
from flask import Flask, render_template, request, jsonify
from dotenv import load_dotenv
from config import Config
from utils.audio import preprocess_to_wav
from utils.summarize import summarize_with_gpt
from faster_whisper import WhisperModel
from datetime import datetime
from flask import send_file
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor

# WARNING NOTIFICATION SKIPPED

warnings.filterwarnings("ignore", message="pkg_resources is deprecated")
warnings.filterwarnings("ignore", message="PySoundFile failed")
warnings.filterwarnings("ignore", message="__audioread_load")
warnings.filterwarnings("ignore", category=UserWarning)
warnings.filterwarnings("ignore", category=FutureWarning)

load_dotenv()


app = Flask(__name__)
app.config.from_object(Config)

def get_supported_compute_type(prefer: str = "int8_float16"):
    """Trả về compute_type khả dụng nhất."""
    has_gpu = torch.cuda.is_available()
    if has_gpu:
        # GPU thường hỗ trợ float16
        return "float16" if prefer in ("int8_float16", "float16") else prefer
    else:
        # CPU thường chỉ hỗ trợ int8 hoặc float32
        return "int8"

def _select_device(device_pref: str):
    if device_pref == "cpu":
        return "cpu"
    if device_pref == "cuda":
        return "cuda"
    try:
        import torch
        return "cuda" if torch.cuda.is_available() else "cpu"
    except Exception:
        return "cpu"

DEVICE = _select_device(app.config["WHISPER_DEVICE"])
model_cache = {}

def get_model(name=None, compute_type=None):
    name = name or app.config["WHISPER_MODEL"]
    compute_type = get_supported_compute_type(compute_type or app.config["WHISPER_COMPUTE"])
    key = (name, DEVICE, compute_type)
    if key not in model_cache:
        model_cache[key] = WhisperModel(
            name,
            device=DEVICE,
            compute_type=compute_type,
        )
    return model_cache[key]

@app.get("/")
def index():
    return render_template(
        "index.html",
        default_model=app.config["WHISPER_MODEL"],
        default_compute=app.config["WHISPER_COMPUTE"],
        default_summary_lang=app.config["DEFAULT_SUMMARY_LANG"],
        device=DEVICE,
    )

def auto_select_params(file_size_mb: float, duration_hint: float = 0):
    """Tự chọn thông số hợp lý."""
    has_gpu = torch.cuda.is_available()
    if file_size_mb < 2:
        model = "base"
        compute = "int8_float16" if not has_gpu else "float16"
        beam = 5
    elif file_size_mb < 10:
        model = "medium"
        compute = "int8_float16" if not has_gpu else "float16"
        beam = 5
    else:
        model = "small"
        compute = "int8"
        beam = 3

    if duration_hint > 300:  # >5 phút
        beam = 3
    return model, compute, beam

@app.post("/api/transcribe")
def api_transcribe():
    try:
        import time
        t0 = time.time()
        file = request.files.get("audio")
        if not file:
            return jsonify({"ok": False, "error": "Missing 'audio' file"}), 400

        # --- AUTO SELECT ---
        file_bytes = file.read()
        size_mb = len(file_bytes) / (1024 * 1024)
        file.seek(0)
        auto_model, auto_compute, auto_beam = auto_select_params(size_mb)

        # ✅ Đảm bảo các biến được gán trước
        model_name = request.form.get("model", "auto")
        if model_name == "auto":
            model_name = auto_model

        compute_type = request.form.get("compute_type", "auto")
        if compute_type == "auto":
            compute_type = auto_compute

        beam_size_str = request.form.get("beam_size", "auto")
        if beam_size_str == "auto":
            beam_size = auto_beam
        else:
            try:
                beam_size = int(beam_size_str)
            except ValueError:
                beam_size = auto_beam

        lang_opt = request.form.get("lang", app.config["DEFAULT_SUMMARY_LANG"])
        asr_lang = request.form.get("asr_lang", "auto")

        # Preprocess (denoise + VAD)
        wav_path, analysis = preprocess_to_wav(file.stream, vad_aggr=app.config["VAD_AGGRESSIVENESS"])

        # Transcribe
        model = get_model(model_name, compute_type)
        print("[INFO] Whisper bắt đầu nhận dạng...")
        segments, info = model.transcribe(
            wav_path,
            vad_filter=False,
            beam_size=beam_size,
            best_of=5,
            patience=0.2,
            temperature=0.0,
            compression_ratio_threshold=2.6,
            log_prob_threshold=-1.0,
            no_speech_threshold=0.45,
            condition_on_previous_text=True,
            language=None if asr_lang == "auto" else asr_lang,
            task="transcribe",
        )
        print("[INFO] Whisper hoàn tất nhận dạng!")


        # Tổng hợp transcript
        text_parts, seg_list = [], []
        for seg in segments:
            seg_list.append({
                "start": round(seg.start, 2),
                "end": round(seg.end, 2),
                "text": seg.text,
                "avg_logprob": getattr(seg, "avg_logprob", None)
            })
            text_parts.append(seg.text)
        transcript = " ".join(text_parts).strip()

        summary = summarize_with_gpt(transcript, lang=lang_opt)
        total_ms = int((time.time() - t0) * 1000)

        return jsonify({
            "ok": True,
            "device": DEVICE,
            "model": model_name,
            "compute_type": compute_type,
            "beam_size": beam_size,
            "analysis": analysis,
            "detected_language": getattr(info, "language", None),
            "duration": getattr(info, "duration", None),
            "processing_ms": total_ms,
            "transcript": transcript,
            "segments": seg_list,
            "summary": summary,
        })

    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

@app.get("/api/export/<fmt>")
def export_recap(fmt):
    """
    Xuất file recap: csv / excel / word
    """
    # Lấy recap JSON mới nhất trong storage
    storage = os.path.join("storage")
    recap_files = sorted(
        [f for f in os.listdir(storage) if f.startswith("recap_") and f.endswith(".json")],
        reverse=True,
    )
    if not recap_files:
        return jsonify({"ok": False, "error": "Không tìm thấy file recap"}), 404

    latest = os.path.join(storage, recap_files[0])
    import json
    with open(latest, encoding="utf-8") as f:
        data = json.load(f)

    transcript = data.get("transcript", "")
    summary = data.get("summary", "")
    segments = data.get("segments", [])

    df = pd.DataFrame(segments)

    if fmt == "csv":
        csv_bytes = df.to_csv(index=False).encode("utf-8")
        return send_file(BytesIO(csv_bytes),
                         mimetype="text/csv",
                         as_attachment=True,
                         download_name="recap.csv")

    elif fmt == "excel":
        output = BytesIO()
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            df.to_excel(writer, sheet_name="Segments", index=False)
            summary_sheet = pd.DataFrame({"Summary": [summary]})
            summary_sheet.to_excel(writer, sheet_name="Summary", index=False)

            ws = writer.sheets["Segments"]
            for col in ws.columns:
                max_len = max(len(str(c.value or "")) for c in col)
                ws.column_dimensions[col[0].column_letter].width = max(12, max_len)
            ws["A1"].font = ws["B1"].font = ws["C1"].font = ws["D1"].font = ws["E1"].font.copy(bold=True)
        output.seek(0)
        return send_file(output,
                         mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                         as_attachment=True,
                         download_name="recap.xlsx")

    elif fmt == "word":
        doc = Document()
        doc.add_heading("RECAP-MAX – Kết quả phân tích", level=1)
        doc.add_paragraph(f"Tóm tắt:\n{summary}", style="Normal")

        doc.add_heading("Chi tiết từng đoạn", level=2)
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "Start"
        hdr[1].text = "End"
        hdr[2].text = "Text"
        hdr[3].text = "Avg LogProb"

        for seg in segments:
            row = table.add_row().cells
            row[0].text = str(seg["start"])
            row[1].text = str(seg["end"])
            row[2].text = seg["text"]
            row[3].text = str(seg.get("avg_logprob", ""))

        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor(30, 30, 30)

        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return send_file(bio,
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         as_attachment=True,
                         download_name="recap.docx")

    return jsonify({"ok": False, "error": "Định dạng không hợp lệ"}), 400

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, threaded=True, debug=True)
